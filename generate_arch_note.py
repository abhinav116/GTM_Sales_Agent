"""Generate Task 1 Architecture Note as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_color(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def add_heading(doc, text, level=1, rgb=(180, 20, 20)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*rgb)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(5)
    return h

def add_body(doc, text, bold=False, italic=False, size=10, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, size=9.5, space_after=3):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def build(path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    p = doc.add_paragraph()
    r = p.add_run('Sales Call Intelligence Agent')
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(180, 20, 20)

    p2 = doc.add_paragraph()
    r2 = p2.add_run('Architecture Note -- Task 1 | Agentic AI Manager Assignment | RAAPID INC')
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(100, 100, 110)
    p2.paragraph_format.space_after = Pt(4)

    p3 = doc.add_paragraph()
    r3 = p3.add_run('Author: Abhinav Marda  |  Stack: Python + Anthropic Claude API  |  June 2026')
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(140, 140, 150)
    p3.paragraph_format.space_after = Pt(12)

    doc.add_paragraph('-' * 90).paragraph_format.space_after = Pt(10)

    # 1. Architecture Overview
    add_heading(doc, '1. Architecture Overview')
    add_body(doc, (
        'The Sales Call Intelligence Agent is a 9-node sequential pipeline that ingests a raw sales '
        'call transcript and produces structured deal intelligence, a CRM update, a follow-up email '
        'draft, and an AE notification. It is designed to run automatically after every sales call, '
        'with a confidence-based human review gate that ensures low-confidence outputs never reach '
        'the CRM without AE validation.'
    ), space_after=10)

    # Pipeline table
    add_body(doc, 'Pipeline Nodes:', bold=True, space_after=6)

    rows = [
        ('Step', 'Node', 'Type', 'Description'),
        ('1', 'Preprocessing', 'Validation', 'Validate transcript length and format. Abort if malformed.'),
        ('2', 'Intelligence Extraction', 'LLM (Claude)', 'Extract pain points, stakeholders, objections, BANT, next steps via enforced JSON schema.'),
        ('3', 'Confidence Scoring', 'Logic', 'Score each extracted field. Flag any field below 0.75 threshold.'),
        ('4', 'Human Review Gate', 'Branch', 'Route flagged fields to AE review queue. Auto-proceed if no flags.'),
        ('5', 'Deal Summary', 'LLM (Claude)', 'Generate CRM-ready deal summary from extracted intelligence.'),
        ('6', 'Email Draft', 'LLM (Claude)', 'Draft follow-up email grounded strictly on transcript content.'),
        ('7', 'Error Handling', 'Retry Logic', 'JSON parse failure triggers correction prompt. Max 3 retries before human escalation.'),
        ('8', 'CRM Push', 'Integration', 'Write structured payload to Salesforce Opportunity via API (mock).'),
        ('9', 'AE Notification', 'Integration', 'Post deal summary to #gtm-deal-intel Slack channel (mock).'),
    ]

    widths = [Cm(1.0), Cm(3.5), Cm(2.8), Cm(8.2)]
    tbl = doc.add_table(rows=len(rows), cols=4)
    tbl.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        for j, (txt, w) in enumerate(zip(row_data, widths)):
            cell = row.cells[j]
            cell.width = w
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(8.5)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_color(cell, 'B41414')
            elif j == 0:
                run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 2. Prompt and Reasoning Logic
    add_heading(doc, '2. Prompt Design & Reasoning Logic')

    add_body(doc, 'System Prompt:', bold=True, space_after=4)
    add_body(doc, (
        'The agent operates under a strict system prompt that establishes three non-negotiable rules: '
        '(1) extract only what is explicitly stated in the transcript, (2) never infer or fabricate, '
        '(3) mark confidence LOW and value null when a field is absent or ambiguous. This is the '
        'primary hallucination prevention mechanism.'
    ), space_after=8)

    add_body(doc, 'Extraction Prompt:', bold=True, space_after=4)
    add_body(doc, (
        'The extraction prompt provides a complete JSON schema with typed fields and confidence score '
        'requirements for every value. Structured output is enforced -- the model is instructed to '
        'return ONLY valid JSON with no commentary. This prevents free text from being passed '
        'downstream, which is a common failure mode in agentic pipelines.'
    ), space_after=8)

    add_body(doc, 'Summary and Email Prompts:', bold=True, space_after=4)
    add_body(doc, (
        'Both downstream prompts receive the extracted JSON as input, not the raw transcript. '
        'This ensures all downstream outputs are grounded on validated, structured intelligence '
        'rather than unprocessed text. The email prompt explicitly instructs the model not to '
        'promise capabilities not discussed in the call.'
    ), space_after=10)

    # 3. Guardrail Design
    add_heading(doc, '3. Guardrail Design')

    guardrails = [
        ('No fact invention', 'System prompt enforces extraction-only behavior. The agent cannot surface a pain point, objection, or stakeholder that is not explicitly present in the transcript.'),
        ('Confidence scoring', 'Every extracted field carries a confidence score (0.0-1.0). Fields below 0.75 are flagged and never written to the CRM without AE confirmation.'),
        ('Human-in-the-loop gate', 'Step 4 routes all flagged fields to the AE review queue before any downstream action. The AE sees exactly which fields are uncertain and why. In the autonomy progression model, this gate is mandatory for the first 90 days.'),
        ('Email send gate', 'The follow-up email draft is delivered to the AE but never auto-sent. The human owns the send button, always.'),
        ('JSON schema enforcement', 'The extraction prompt specifies an exact JSON schema. Parse failures trigger a correction prompt (up to 3 retries). If all retries fail, the run halts and is escalated to human review -- no partial data is written to the CRM.'),
        ('Scope separation', 'The agent has no access to PHI, clinical records, or the core RAAPID product. It operates exclusively on sales call transcripts.'),
    ]

    for title, desc in guardrails:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{title}: ')
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 4. Failure Modes and Limitations
    add_heading(doc, '4. Failure Modes & Limitations')

    failures = [
        ('JSON parse failure', 'Handled via retry loop with correction prompt. Max 3 attempts. Demonstrated in Step 7.'),
        ('Low-confidence extraction', 'Handled via confidence gate. Flagged fields go to human review queue.'),
        ('Empty or malformed transcript', 'Caught in Step 1 preprocessing. Pipeline aborts cleanly with no CRM write.'),
        ('LLM hallucination', 'Mitigated by system prompt grounding rules and downstream human review gate. Not fully eliminable -- spot-check evals required.'),
        ('API timeout / rate limit', 'Not implemented in demo. Production version would add exponential backoff with a max wait of 30s before human escalation.'),
        ('Scope drift', 'Agent prompt does not prevent a user from feeding non-sales content as a transcript. Input validation (topic classification) would be added in production.'),
    ]

    fail_widths = [Cm(4.0), Cm(11.5)]
    fail_tbl = doc.add_table(rows=len(failures) + 1, cols=2)
    fail_tbl.style = 'Table Grid'

    header_row = fail_tbl.rows[0]
    for j, (txt, w) in enumerate(zip(['Failure Mode', 'Handling'], fail_widths)):
        cell = header_row.cells[j]
        cell.width = w
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_color(cell, 'B41414')

    for i, (mode, handling) in enumerate(failures):
        row = fail_tbl.rows[i + 1]
        for j, (txt, w) in enumerate(zip([mode, handling], fail_widths)):
            cell = row.cells[j]
            cell.width = w
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(8.5)
            if j == 0:
                run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 5. Productionization inside RAAPID
    add_heading(doc, '5. How to Productionize & Govern Inside RAAPID')

    add_body(doc, 'Phase 1 -- Shadow Mode (weeks 1-4):', bold=True, space_after=4)
    add_bullet(doc, 'Run agent in parallel with existing manual post-call process.')
    add_bullet(doc, 'AE compares agent output to their own notes. No CRM writes yet.')
    add_bullet(doc, 'Track extraction accuracy and escalation rate as baseline metrics.')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_body(doc, 'Phase 2 -- Assisted (weeks 4-16):', bold=True, space_after=4)
    add_bullet(doc, 'Agent outputs used as CRM draft -- AE reviews and confirms before write.')
    add_bullet(doc, 'All flagged fields require explicit AE approval.')
    add_bullet(doc, 'Target: escalation rate below 15% by week 8.')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_body(doc, 'Phase 3 -- Co-pilot (months 4-9):', bold=True, space_after=4)
    add_bullet(doc, 'High-confidence fields (>0.75) auto-write to CRM.')
    add_bullet(doc, 'AE only reviews flagged exceptions.')
    add_bullet(doc, 'Gate: escalation rate < 10% AND AE satisfaction > 4/5.')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_body(doc, 'Phase 4 -- Autonomous (month 9+):', bold=True, space_after=4)
    add_bullet(doc, 'Full CRM update after every call. Email draft delivered to AE inbox for one-click approve.')
    add_bullet(doc, 'Human only reviews exceptions surfaced by the agent.')

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_body(doc, 'Governance requirements:', bold=True, space_after=4)
    add_bullet(doc, 'Immutable audit log for every agent run: transcript hash, extracted output, confidence scores, human review outcome.')
    add_bullet(doc, 'No PHI in any agent input or output. Transcripts are sales conversations only.')
    add_bullet(doc, 'Weekly review of flagged runs with GTM lead for first 90 days.')
    add_bullet(doc, 'Model version pinned. Any Claude model update requires re-validation of extraction accuracy before rollout.')
    add_bullet(doc, 'Quarterly review with legal to confirm no agent scope drift toward regulated activity.')

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 6. Cost and Scaling
    add_heading(doc, '6. Cost, Latency & Scaling')

    metrics = [
        ('Tokens per run', '~2,400 (3 LLM calls: extraction + summary + email)'),
        ('Cost per run', '~$0.012 at claude-sonnet-4-6 pricing'),
        ('50 calls/day', '~$18/month'),
        ('200 calls/day', '~$72/month'),
        ('End-to-end latency', '~4-6 seconds (live API)'),
        ('Scaling bottleneck', 'Claude API rate limits at high volume -- mitigated by async batch processing'),
        ('Storage', 'CRM payload JSON per run (~3KB). 200 runs/day = ~600KB/day, negligible'),
    ]

    m_widths = [Cm(5.0), Cm(10.5)]
    m_tbl = doc.add_table(rows=len(metrics) + 1, cols=2)
    m_tbl.style = 'Table Grid'

    hrow = m_tbl.rows[0]
    for j, (txt, w) in enumerate(zip(['Metric', 'Value'], m_widths)):
        cell = hrow.cells[j]
        cell.width = w
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_color(cell, 'B41414')

    for i, (metric, val) in enumerate(metrics):
        row = m_tbl.rows[i + 1]
        for j, (txt, w) in enumerate(zip([metric, val], m_widths)):
            cell = row.cells[j]
            cell.width = w
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(8.5)
            if j == 0:
                run.font.bold = True

    doc.save(path)
    print(f'Architecture note saved: {path}')

if __name__ == '__main__':
    out = os.path.join(os.path.expanduser('~'), 'Downloads', 'raapid_agent', 'RAAPID_Task1_Architecture_Note.docx')
    build(out)
